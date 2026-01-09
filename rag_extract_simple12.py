"""
rag_extract_simple.py

Extracts actions with RAG status from Technician Commitment documents.
Handles RAG detection through explicit R/A/G labels and AI-based contextual identification.

Outputs data in format ready for manual/automated alignment workflows:
- First 5 columns populated: LineID, Institution, AP/RAG_AP Number, AP, RAG_AP
- Remaining columns added as headers only (for later population)

Author: Simplified from rag_extract_longitudinal.py
"""

import csv
import io
import os
import re

import pdfplumber

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX support will be limited.")

# Vertex AI for action extraction
import vertexai
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
from vertexai.generative_models import GenerationConfig, GenerativeModel

# --- CONFIGURATION ---
SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
CREDENTIALS_FILE = "credentials.json"
TOKEN_FILE = "token.json"
LOCATION = "us-central1"
OUTPUT_CSV_FILE = "rag_actions_extraction.csv"

PROJECT_ID = "tcdocprog"  # Update as needed
MODEL_ID = os.environ.get("GCP_MODEL_ID", "gemini-2.0-flash")

# Filename parsing patterns to extract institution, year, and document type
FILENAME_PATTERNS = [
    r"(?P<year>\d{4})-(?P<institution>[A-Za-z]+)-(?P<doctype>.*)",
    r"(?P<institution>[A-Za-z]+)-(?P<year>\d{4})-(?P<doctype>.*)",
    r"(?P<institution>[A-Za-z]+)(?P<doctype>ActionPlan|RAG)(?P<year>\d{4})",
]


def authenticate():
    """Handles OAuth2 authentication and returns valid credentials."""
    creds = None
    if os.path.exists(TOKEN_FILE):
        creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists(CREDENTIALS_FILE):
                print(f"ERROR: Credentials file '{CREDENTIALS_FILE}' not found.")
                return None
            flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_FILE, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(TOKEN_FILE, "w") as token:
            token.write(creds.to_json())
    return creds


def parse_filename_metadata(filename):
    """
    Extract institution, year, and document type from filename.
    Returns dict with keys: institution, year, doc_type, is_rag
    """
    metadata = {
        "institution": "Unknown",
        "year": None,
        "doc_type": "Unknown",
        "is_rag": False,
    }

    # Try each pattern
    for pattern in FILENAME_PATTERNS:
        match = re.search(pattern, filename, re.IGNORECASE)
        if match:
            groups = match.groupdict()
            metadata["institution"] = groups.get("institution", "Unknown")
            metadata["year"] = groups.get("year", None)
            metadata["doc_type"] = groups.get("doctype", "Unknown")
            break

    # Check if it's a RAG document
    metadata["is_rag"] = bool(re.search(r"RAG|rag", filename))

    return metadata


def parse_stage_number(filename):
    """
    Extract stage number from filename (1, 2, or 3).
    Returns integer or None if not found.
    """
    filename_upper = filename.upper()

    # Pattern 1: Combined RAG documents (e.g., RAG_ActionPlan2-ActionPlan3)
    combined_match = re.search(r"RAG_?ACTIONPLAN(\d+)-ACTIONPLAN(\d+)", filename_upper)
    if combined_match:
        # For combined, use the first number
        return int(combined_match.group(1))

    # Pattern 2: Single RAG document (e.g., RAG_ActionPlan1, RAG_ActionPlan2)
    rag_match = re.search(r"RAG_?ACTIONPLAN(\d+)", filename_upper)
    if rag_match:
        return int(rag_match.group(1))

    # Pattern 3: Inst3-style RAG (e.g., RAG1, RAG2)
    Inst3_rag_simple = re.search(r"RAG(\d+)(?!.*ACTIONPLAN)", filename_upper)
    if Inst3_rag_simple:
        return int(Inst3_rag_simple.group(1))

    # Pattern 3b: RAG-AP format (e.g., RAG-AP1)
    Inst3_rag_dash = re.search(r"RAG-AP(\d+)", filename_upper)
    if Inst3_rag_dash:
        return int(Inst3_rag_dash.group(1))

    # Pattern 4: Inst3-style Action Plans (e.g., AP1, AP2, Inst3-2024-AP1)
    Inst3_ap_match = re.search(r"AP(\d+)", filename_upper)
    if Inst3_ap_match:
        return int(Inst3_ap_match.group(1))

    # Pattern 5: Regular Action Plan (e.g., ActionPlan1, ActionPlan2)
    ap_match = re.search(r"ACTIONPLAN(\d+)", filename_upper)
    if ap_match:
        return int(ap_match.group(1))

    return None


def extract_explicit_rag(text):
    """
    Extract RAG status from text using explicit markers.
    Returns 'red', 'amber', 'green', or None
    """
    text_lower = text.lower()

    # Look for explicit RAG markers
    rag_patterns = [
        (r"\brag[:\s]*red\b", "red"),
        (r"\brag[:\s]*r\b", "red"),
        (r"\bred\b", "red"),
        (r"\brag[:\s]*amber\b", "amber"),
        (r"\brag[:\s]*a\b", "amber"),
        (r"\bamber\b", "amber"),
        (r"\brag[:\s]*green\b", "green"),
        (r"\brag[:\s]*g\b", "green"),
        (r"\bgreen\b", "green"),
    ]

    for pattern, status in rag_patterns:
        if re.search(pattern, text_lower):
            return status

    return None


def get_files_from_folder(service, folder_id):
    """Retrieve list of PDF/DOCX/Google Docs files from Google Drive folder."""
    query = f"'{folder_id}' in parents and (mimeType='application/pdf' or mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or mimeType='application/vnd.google-apps.document')"
    try:
        results = (
            service.files().list(q=query, fields="files(id, name, mimeType)").execute()
        )
        return results.get("files", [])
    except HttpError as error:
        print(f"An error occurred: {error}")
        return []


def download_file(service, file_id, mime_type):
    """Download file from Google Drive and return as bytes.
    For Google Docs, exports to DOCX format to preserve structure."""
    try:
        # Google Docs need to be exported, not downloaded
        if mime_type == "application/vnd.google-apps.document":
            request = service.files().export_media(
                fileId=file_id,
                mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            request = service.files().get_media(fileId=file_id)

        file_stream = io.BytesIO()
        downloader = MediaIoBaseDownload(file_stream, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        file_stream.seek(0)
        return file_stream.read()
    except HttpError as error:
        print(f"An error occurred: {error}")
        return None


def extract_text_from_pdf(pdf_bytes):
    """Extract text from PDF bytes.
    Returns tuple: (text_string, empty_list) to match DOCX function signature.
    """
    text = ""
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text, []  # Return empty list for structured_actions


def extract_text_from_docx(docx_bytes):
    """Extract text from DOCX bytes, handling both paragraphs and tables.
    Returns tuple: (text_string, list_of_action_rag_pairs)

    For tables that appear to be action tables (with RAG column), extracts
    structured data with actions and their RAG status.
    Handles continuation tables and merged header rows.
    """
    if not DOCX_AVAILABLE:
        return "", []

    text_parts = []
    action_rag_pairs = []
    doc = Document(io.BytesIO(docx_bytes))

    # Extract from paragraphs
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if para_text:
            text_parts.append(para_text)

    # Track column positions across tables (for continuation tables)
    last_action_col = None
    last_rag_col = None

    # Extract from tables
    for table_idx, table in enumerate(doc.tables):
        rag_column_idx = None
        action_column_idx = None
        start_row = 1  # Usually skip row 0 (header)

        if len(table.rows) == 0:
            continue

        # Check row 0 for headers
        header_row = table.rows[0]
        potential_action_cols = []  # Track all potential action columns

        # First check if row 0 is a merged title row (all cells have same text)
        row0_texts = [cell.text.strip() for cell in header_row.cells]
        is_merged_title = (
            len(set(row0_texts)) == 1 and row0_texts[0]
        )  # All same and non-empty

        if not is_merged_title:
            # Row 0 is actual headers, check for action columns
            for idx, cell in enumerate(header_row.cells):
                cell_text = cell.text.strip().lower()
                if "rag" == cell_text or "rag" in cell_text.split():
                    rag_column_idx = idx

                # Check for action columns, but exclude columns that are clearly not actions
                if (
                    "action" in cell_text
                    or "step" in cell_text
                    or "what is to be done" in cell_text
                ):
                    # Only apply exclusions to avoid non-action columns
                    if not any(
                        exclude in cell_text
                        for exclude in [
                            "next",
                            "outcome",
                            "impact",
                            "achieve",
                            "result",
                            "success",
                        ]
                    ):
                        potential_action_cols.append(idx)

        # If row 0 was merged title OR didn't find action columns, check row 1
        if (is_merged_title or len(potential_action_cols) == 0) and len(table.rows) > 1:
            second_row = table.rows[1]
            potential_action_cols = []  # Reset for row 1

            for idx, cell in enumerate(second_row.cells):
                cell_text = cell.text.strip().lower()
                if "rag" == cell_text or "rag" in cell_text.split():
                    rag_column_idx = idx
                    start_row = 2  # Data starts from row 2

                # Check for action columns with exclusions
                if (
                    "action" in cell_text
                    or "step" in cell_text
                    or "what is to be done" in cell_text
                ):
                    if not any(
                        exclude in cell_text
                        for exclude in [
                            "next",
                            "outcome",
                            "impact",
                            "achieve",
                            "result",
                            "success",
                        ]
                    ):
                        potential_action_cols.append(idx)
                        start_row = 2  # Set start_row when we find action columns

        # If multiple columns have "action" in header, choose the best one
        # Prioritize: 1) Exact "action" match, 2) Most text in first data row
        if len(potential_action_cols) > 1 and len(table.rows) > start_row:
            # First, prefer columns with exact "action" match
            exact_action_cols = []
            check_row = table.rows[1] if start_row == 2 else table.rows[0]

            for col_idx in potential_action_cols:
                header_text = check_row.cells[col_idx].text.strip().lower()
                if header_text in ["action", "what is to be done"]:
                    exact_action_cols.append(col_idx)

            # If we have exact "action" matches, choose from those
            candidates = (
                exact_action_cols if exact_action_cols else potential_action_cols
            )

            # Among candidates, pick the one with most text in first NON-EMPTY data row
            best_col = candidates[0]
            max_text_len = 0

            # Find first non-empty row starting from start_row
            for data_row_idx in range(start_row, min(start_row + 5, len(table.rows))):
                # Check if this row has any substantial content
                has_content = False
                for col_idx in candidates:
                    if len(table.rows[data_row_idx].cells) > col_idx:
                        text_len = len(
                            table.rows[data_row_idx].cells[col_idx].text.strip()
                        )
                        if text_len > 5:  # At least some content
                            has_content = True
                            break

                if has_content:
                    # This row has content, use it to decide
                    for col_idx in candidates:
                        if len(table.rows[data_row_idx].cells) > col_idx:
                            text_len = len(
                                table.rows[data_row_idx].cells[col_idx].text.strip()
                            )
                            if text_len > max_text_len:
                                max_text_len = text_len
                                best_col = col_idx
                    break  # Found a good row, stop searching

            action_column_idx = best_col
        elif len(potential_action_cols) == 1:
            action_column_idx = potential_action_cols[0]

        # If no headers found in this table, check if it's a continuation table
        # (use column positions from previous table)
        is_continuation = False
        if (rag_column_idx is None or action_column_idx is None) and (
            last_rag_col is not None and last_action_col is not None
        ):
            rag_column_idx = last_rag_col
            action_column_idx = last_action_col
            start_row = 0  # No header in continuation tables
            is_continuation = True

        # If we found action and RAG columns, extract data
        if rag_column_idx is not None and action_column_idx is not None:
            if not is_continuation:
                print(
                    f"  DEBUG: Table {table_idx} - Found action column at {action_column_idx}, RAG column at {rag_column_idx}, start_row={start_row}"
                )
                # Remember these positions for potential continuation tables
                last_action_col = action_column_idx
                last_rag_col = rag_column_idx
            else:
                print(
                    f"  DEBUG: Table {table_idx} - Continuation table using columns: action={action_column_idx}, RAG={rag_column_idx}"
                )

            for row_idx in range(start_row, len(table.rows)):
                row = table.rows[row_idx]

                # Check we have enough columns
                if len(row.cells) <= max(rag_column_idx, action_column_idx):
                    continue

                action_cell = row.cells[action_column_idx].text.strip()
                rag_cell = row.cells[rag_column_idx].text.strip()

                # Clean up action text (remove internal line breaks)
                action_text = " ".join(action_cell.split("\n")).strip()

                # Only process if there's actual action text (and not just a header remnant)
                if not action_text or len(action_text) < 5:
                    continue

                # Skip if this looks like a header row
                if action_text.lower() in [
                    "action",
                    "actions",
                    "steps",
                    "what is to be done",
                ]:
                    continue

                # Extract RAG status with robust detection
                rag_status = None
                rag_lower = rag_cell.lower().strip()

                if "green" in rag_lower or rag_lower == "g":
                    rag_status = "green"
                elif "amber" in rag_lower or rag_lower == "a":
                    rag_status = "amber"
                elif "red" in rag_lower or rag_lower == "r":
                    rag_status = "red"

                # Debug output
                if rag_status:
                    print(f"    Row {row_idx}: Found RAG={rag_status}")
                else:
                    print(
                        f"    Row {row_idx}: WARNING - No RAG found in cell: '{rag_cell[:30]}'"
                    )

                action_rag_pairs.append(
                    {"action": action_text, "rag_status": rag_status}
                )
        else:
            # No structured RAG table, just extract text
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cell_text = " ".join(cell_text.split("\n"))
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

    text = "\n".join(text_parts)

    if action_rag_pairs:
        print(f"  DEBUG: Total extracted pairs: {len(action_rag_pairs)}")

        # Merge continuation rows (actions split across multiple rows)
        merged_pairs = []
        i = 0
        while i < len(action_rag_pairs):
            current = action_rag_pairs[i]
            current_action = current["action"]

            # Check if next row is a continuation (starts with lowercase or common continuation words)
            # and current action seems incomplete (doesn't end with period or other terminal punctuation)
            while i + 1 < len(action_rag_pairs):
                next_pair = action_rag_pairs[i + 1]
                next_action = next_pair["action"]

                # Signs this might be a continuation:
                # 1. Next action starts with lowercase letter
                # 2. Current action doesn't end with proper punctuation
                # 3. Next action doesn't start with typical action markers (numbers, bullets, capitals)
                is_continuation = (
                    next_action
                    and len(next_action) > 0
                    and next_action[0].islower()  # Starts with lowercase
                    and not current_action.rstrip().endswith(
                        (".", "!", "?", ")", "]")
                    )  # Current doesn't end properly
                    and not next_action[0].isdigit()  # Next doesn't start with number
                    and not next_action.startswith(
                        "("
                    )  # Next doesn't start with parenthesis (like "(i)")
                )

                if is_continuation:
                    # Merge the continuation into current action
                    current_action = current_action.rstrip() + " " + next_action
                    # Keep the RAG status from the first (main) action, not the continuation
                    i += 1  # Skip the continuation row
                else:
                    break

            # Add the merged action
            merged_pairs.append(
                {
                    "action": current_action,
                    "rag_status": current["rag_status"],  # Use RAG from first part
                }
            )
            i += 1

        action_rag_pairs = merged_pairs
        print(f"  DEBUG: After merging continuations: {len(action_rag_pairs)}")
        rag_count = sum(1 for p in action_rag_pairs if p["rag_status"])
        print(f"  DEBUG: Pairs with RAG status: {rag_count}")

    return text, action_rag_pairs


def extract_actions_with_gemini(text, file_name):
    """
    Use Gemini to extract actions from document text.
    Returns list of action strings.
    """
    vertexai.init(project=PROJECT_ID, location=LOCATION)
    model = GenerativeModel(MODEL_ID)

    prompt = f"""You are analyzing a Technician Commitment action plan document.

Extract all discrete actions/commitments from this document. Each action should be:
- A complete, coherent sentence or set of related sentences describing one action
- NOT split mid-sentence or mid-phrase
- Distinct from other actions

IMPORTANT RULES:
1. Return ONLY the action text itself - do NOT include any numbering, bullets, or prefixes
2. REMOVE any RAG status indicators (Red, Amber, Green, R, A, G) from the action text
3. Ensure each action is a COMPLETE sentence - do not split sentences across multiple lines
4. If an action spans multiple lines in the original, combine it into one continuous text
5. Clean up any formatting artifacts (extra spaces, line breaks within sentences)

For each action, return it on a new line.
Do not include headers, table labels, section titles, or explanatory text.

Document text:
{text[:50000]}

Return the actions, one per line, with RAG status removed:"""

    generation_config = GenerationConfig(
        temperature=0.1,
        top_p=0.95,
        max_output_tokens=8192,
    )

    try:
        response = model.generate_content(
            prompt,
            generation_config=generation_config,
        )

        if not response.text:
            print(f"  ! No response text from Gemini for {file_name}")
            return []

        # Split response into individual actions
        actions = [
            line.strip() for line in response.text.strip().split("\n") if line.strip()
        ]
        return actions

    except Exception as e:
        print(f"  ! Error calling Gemini for {file_name}: {e}")
        return []


def process_file_for_extraction(service, file_item):
    """
    Process a single file and extract actions.
    Returns list of dicts with action data.
    """
    file_id = file_item["id"]
    file_name = file_item["name"]
    mime_type = file_item["mimeType"]

    print(f"\nProcessing: {file_name}")

    # Parse metadata
    metadata = parse_filename_metadata(file_name)
    stage_number = parse_stage_number(file_name)

    if stage_number is None:
        print("  ! Could not determine stage number from filename")
        return []

    # Download file
    file_bytes = download_file(service, file_id, mime_type)
    if not file_bytes:
        print(f"  ! Failed to download {file_name}")
        return []

    # Extract text
    text = ""
    structured_actions = []

    if mime_type == "application/pdf":
        text, structured_actions = extract_text_from_pdf(file_bytes)
    elif mime_type == "application/vnd.google-apps.document":
        # Google Docs are exported as DOCX, use DOCX extraction
        text, structured_actions = extract_text_from_docx(file_bytes)
    elif (
        mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        # Native DOCX files
        text, structured_actions = extract_text_from_docx(file_bytes)
    else:
        print(f"  ! Unsupported mime type: {mime_type}")
        return []

    # If we have structured actions from table extraction, use those directly
    if structured_actions:
        print(f"  > Found {len(structured_actions)} actions in structured table format")
        results = []
        for item in structured_actions:
            results.append(
                {
                    "institution": metadata["institution"],
                    "stage_number": stage_number,
                    "is_rag": metadata["is_rag"],
                    "action": item["action"],
                    "rag_status": item["rag_status"],
                }
            )
        print(f"  > Extracted {len(results)} actions from tables")
        return results

    if not text.strip():
        print(f"  ! No text extracted from {file_name}")
        return []

    # Extract actions using Gemini
    actions = extract_actions_with_gemini(text, file_name)

    if not actions:
        print(f"  ! No actions extracted from {file_name}")
        return []

    # Process each action
    results = []
    for action_text in actions:
        # Clean up the action text
        # Remove leading/trailing whitespace
        action_text = action_text.strip()

        # Skip empty actions
        if not action_text:
            continue

        # Extract explicit RAG from action text BEFORE cleaning
        rag_explicit = extract_explicit_rag(action_text)

        # Remove RAG status words from the action text itself
        # This prevents "Green", "Amber", "Red" from appearing in the action
        action_cleaned = action_text
        for rag_word in [
            "Green",
            "Amber",
            "Red",
            "green",
            "amber",
            "red",
            "RAG: Green",
            "RAG: Amber",
            "RAG: Red",
            "RAG:Green",
            "RAG:Amber",
            "RAG:Red",
            "RAG Green",
            "RAG Amber",
            "RAG Red",
        ]:
            # Remove the RAG word if it appears at the end or as a standalone element
            action_cleaned = re.sub(
                r"\s*\|\s*" + re.escape(rag_word) + r"\s*$", "", action_cleaned
            )
            action_cleaned = re.sub(
                r"\s+" + re.escape(rag_word) + r"\s*$", "", action_cleaned
            )
            action_cleaned = re.sub(
                r"^\s*" + re.escape(rag_word) + r"\s*\|", "", action_cleaned
            )

        # Clean up any remaining artifacts
        action_cleaned = re.sub(r"\s+", " ", action_cleaned).strip()
        action_cleaned = re.sub(r"\|\s*$", "", action_cleaned).strip()
        action_cleaned = re.sub(r"^\s*\|", "", action_cleaned).strip()

        # Skip if action becomes empty after cleaning
        if not action_cleaned:
            continue

        # Try AI-based RAG identification for RAG documents
        rag_from_ai = None
        if metadata["is_rag"] and not rag_explicit:
            # Could add AI-based RAG detection here if needed
            pass

        # Priority: explicit > AI-identified > None
        rag_status = rag_explicit or rag_from_ai

        results.append(
            {
                "institution": metadata["institution"],
                "stage_number": stage_number,
                "is_rag": metadata["is_rag"],
                "action": action_cleaned,
                "rag_status": rag_status,
            }
        )

    print(f"  > Extracted {len(results)} actions")
    return results


def main():
    """Main execution function."""
    creds = authenticate()
    if not creds:
        print("Authentication failed. Exiting.")
        return

    try:
        service = build("drive", "v3", credentials=creds)
    except Exception as e:
        print(f"Could not build Google Drive service: {e}")
        return

    # Get folder ID
    folder_id = input(
        "Enter Google Drive folder ID (or press Enter for default): "
    ).strip()
    if not folder_id:
        folder_id = "1Gm5spunt70hOLj-YDQyTmMk_Kr3HvUK9"  # Default

    print(f"\nUsing model '{MODEL_ID}' in project '{PROJECT_ID}'...")
    print(f"Output will be saved to '{OUTPUT_CSV_FILE}'\n")

    files = get_files_from_folder(service, folder_id)
    print(f"Found {len(files)} files.\n")

    all_results = []

    for file_item in files:
        file_results = process_file_for_extraction(service, file_item)
        all_results.extend(file_results)

    # Group by institution and stage for organization
    print("\n" + "=" * 60)
    print("ORGANIZING EXTRACTED ACTIONS...")
    print("=" * 60 + "\n")

    # Organize actions by institution, stage, and type
    organized = {}
    for result in all_results:
        institution = result["institution"]
        stage_num = result["stage_number"]
        is_rag = result["is_rag"]

        key = (institution, stage_num, is_rag)
        if key not in organized:
            organized[key] = []
        organized[key].append(result)

    # Create output rows in sequential format
    output_rows = []
    line_id = 1

    # Sort by institution, then stage number, then document type (AP before RAG_AP)
    for (institution, stage_num, is_rag), actions in sorted(
        organized.items(), key=lambda x: (x[0][0], x[0][1], x[0][2])
    ):
        # Create document type label
        doc_type = f"RAG_AP{stage_num}" if is_rag else f"AP{stage_num}"

        # Add each action as a separate row
        for action_data in actions:
            action_text = action_data["action"]
            rag_value = action_data.get("rag_status", "") if is_rag else ""

            # Convert rag_status to single letter format
            if rag_value:
                rag_value = rag_value[0].upper()  # R, A, or G

            output_rows.append(
                {
                    "LineID": line_id,
                    "Institution": institution,
                    "Document Type": doc_type,
                    "Action": action_text,
                    "RAG value": rag_value,
                    "Category": "",
                }
            )

            line_id += 1

    # Write to CSV
    if output_rows:
        fieldnames = [
            "LineID",
            "Institution",
            "Document Type",
            "Action",
            "RAG value",
            "Category",
        ]

        with open(OUTPUT_CSV_FILE, "w", newline="", encoding="utf-8") as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(output_rows)

        print(f"\n{'=' * 60}")
        print(f"SUCCESS: Wrote {len(output_rows)} rows to '{OUTPUT_CSV_FILE}'")
        print(f"{'=' * 60}\n")

        # Summary statistics
        institutions = set(row["Institution"] for row in output_rows)
        doc_types = set(row["Document Type"] for row in output_rows)

        print("Extraction Summary:")
        print(f"  Institutions: {', '.join(sorted(institutions))}")
        print(f"  Document types: {', '.join(sorted(doc_types))}")
        print(f"  Total action rows: {len(output_rows)}")

        # RAG status distribution
        rag_counts = {"R": 0, "A": 0, "G": 0, "none": 0}
        for row in output_rows:
            status = row["RAG value"] or "none"
            rag_counts[status] = rag_counts.get(status, 0) + 1

        print("\nRAG Status Distribution:")
        for status, count in rag_counts.items():
            print(f"  {status}: {count}")

        # Count by document type
        doc_type_counts = {}
        for row in output_rows:
            doc_type = row["Document Type"]
            doc_type_counts[doc_type] = doc_type_counts.get(doc_type, 0) + 1

        print("\nActions by Document Type:")
        for doc_type in sorted(doc_type_counts.keys()):
            print(f"  {doc_type}: {doc_type_counts[doc_type]} actions")
    else:
        print("\nNo actions were extracted.")


if __name__ == "__main__":
    main()
